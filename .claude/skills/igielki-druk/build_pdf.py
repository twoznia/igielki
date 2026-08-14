#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Buduje wersję do druku (PDF/DOCX) powieści IGIEŁKI.

Hierarchia: TOM → CZĘŚĆ → ROZDZIAŁ (widoczna w spisie treści i w zakładkach PDF).
Źródłem jest struktura katalogów `opowiadania/` utworzona przez `split_source.py`.

Domyślnie (bez argumentów): PDF **całości** (wszystkie cztery tomy).

Parametry:
  --typ  pdf | docx | oba          (domyślnie: pdf)
  --tom  N | calosc | wszystko     (domyślnie: calosc)
       N        -> jeden tom (np. --tom 2)
       calosc   -> jeden dokument z całą serią (tomy + części)
       wszystko -> każdy tom osobno + całość
  --out ŚCIEŻKA                    (tylko dla pojedynczego celu, bez rozszerzenia)

WAŻNE: PDF powstaje CZYSTO w Pythonie (reportlab) — bez MS Word i bez drukarki.
Cechy PDF:
  - klikalny **spis treści** (Tom → Część → Rozdział) z numerami stron,
  - **zakładki/outline po lewej**: Tom → Część → Rozdział, klikalne,
  - każdy tom, część i rozdział od nowej strony; numeracja stron,
  - separator sceny `---` renderowany jako odstęp z gwiazdką,
  - czcionka Montserrat, jeśli TTF jest w `fonts/`/systemie; inaczej Arial.
"""
import argparse
import re
import sys
import os
from pathlib import Path

FONT = "Montserrat"
SIZE_BODY = 12
SIZE_CHAPTER = 15     # rozdział
SIZE_PART = 20        # część
SIZE_TOM = 26         # tom
SIZE_SERIES = 14

SERIES_TITLE = "IGIEŁKI"
SERIES_SUB = "Wielka Księga Mchowa. Wszystkie tomy"
SERIES_BLURB = ("Rodzina jeży z Mchowa Dolnego. Cztery tomy o tym, że żeby żyć "
                "z kimś obok, trzeba nauczyć się trzech rzeczy: pytać, dzielić i myć zęby.")

REPO = Path(__file__).resolve().parents[3]
STORIES_DIR = REPO / "opowiadania"
OUT_DIR = REPO / "druk"

INLINE = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*")
ORDINAL = ["", "pierwszy", "drugi", "trzeci", "czwarty", "piąty", "szósty",
           "siódmy", "ósmy", "dziewiąty", "dziesiąty"]


# ----------------------------- struktura -----------------------------

def _title_of(md_path):
    """Pierwszy nagłówek '# ' w pliku (albo nazwa pliku)."""
    if md_path.exists():
        for line in md_path.read_text(encoding="utf-8").splitlines():
            m = re.match(r"^#\s+(.*)$", line)
            if m:
                return m.group(1).strip()
    return md_path.stem


def list_toms():
    return sorted([d for d in STORIES_DIR.glob("tom-*") if d.is_dir()],
                  key=lambda p: p.name)


def tom_number(tom_dir):
    m = re.match(r"tom-(\d+)", tom_dir.name)
    return int(m.group(1)) if m else 0


def parts_of(tom_dir):
    return sorted([d for d in tom_dir.glob("czesc-*") if d.is_dir()],
                  key=lambda p: p.name)


def chapters_of(part_dir):
    return sorted([p for p in part_dir.glob("*.md") if not p.name.startswith("_")],
                  key=lambda p: p.name)


def chapter_no(md_path):
    m = re.match(r"^\s*(\d+)", md_path.name)
    return int(m.group(1)) if m else 0


def body_lines(md_path):
    """Zwraca listę ('chapter'|'scene'|'body'|'quote', tekst).

    Pierwszy '# ' (tytuł rozdziału) pomijamy — tytuł składany jest osobno.
    """
    out, seen_title = [], False
    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("# ") and not seen_title:
            seen_title = True
            continue
        if line.strip() == "---":
            out.append(("scene", ""))
        elif line.startswith("> "):
            out.append(("quote", line[2:].strip()))
        else:
            out.append(("body", line))
    return out


# ============================ PDF (reportlab) ============================

def _find_font_family():
    dirs = [REPO / "fonts", Path(os.environ.get("WINDIR", "C:/Windows")) / "Fonts"]

    def find(names):
        for d in dirs:
            for n in names:
                p = d / n
                if p.exists():
                    return str(p)
        return None
    mont = (find(["Montserrat-Regular.ttf", "Montserrat.ttf"]),
            find(["Montserrat-Bold.ttf", "Montserrat-SemiBold.ttf"]),
            find(["Montserrat-Italic.ttf"]),
            find(["Montserrat-BoldItalic.ttf"]))
    if mont[0] and mont[1]:
        return mont[0], mont[1], mont[2] or mont[0], mont[3] or mont[1], "Montserrat"
    ar = (find(["arial.ttf", "Arial.ttf"]), find(["arialbd.ttf"]),
          find(["ariali.ttf"]), find(["arialbi.ttf"]))
    if ar[0] and ar[1]:
        return ar[0], ar[1], ar[2] or ar[0], ar[3] or ar[1], "Arial"
    return None


def _register_fonts():
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    fam = _find_font_family()
    if not fam:
        print("UWAGA: brak Montserrat/Arial TTF — PDF użyje Helvetica.", file=sys.stderr)
        return ("Helvetica", "Helvetica-Bold", "Helvetica-Oblique",
                "Helvetica-BoldOblique", "Helvetica")
    reg, bold, ital, bi, label = fam
    base = "IG"
    pdfmetrics.registerFont(TTFont(base, reg))
    pdfmetrics.registerFont(TTFont(base + "-B", bold))
    pdfmetrics.registerFont(TTFont(base + "-I", ital))
    pdfmetrics.registerFont(TTFont(base + "-BI", bi))
    pdfmetrics.registerFontFamily(base, normal=base, bold=base + "-B",
                                  italic=base + "-I", boldItalic=base + "-BI")
    if label != "Montserrat":
        print(f"Info: Montserrat nie znaleziony — PDF użyje czcionki {label}.", file=sys.stderr)
    return base, base + "-B", base + "-I", base + "-BI", label


def _rml(text):
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
    return text


def build_pdf(targets):
    from reportlab.platypus import (BaseDocTemplate, PageTemplate, Frame, Paragraph,
                                    Spacer, PageBreak)
    from reportlab.platypus.tableofcontents import TableOfContents
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT

    REG, BOLD, ITAL, BI, label = _register_fonts()

    body = ParagraphStyle("body", fontName=REG, fontSize=SIZE_BODY, leading=SIZE_BODY * 1.4,
                          alignment=TA_JUSTIFY, spaceAfter=6, firstLineIndent=0.6 * cm)
    quote = ParagraphStyle("quote", parent=body, leftIndent=0.6 * cm, rightIndent=0.6 * cm,
                           firstLineIndent=0, fontName=ITAL)
    scene = ParagraphStyle("scene", fontName=REG, fontSize=SIZE_BODY, alignment=TA_CENTER,
                           spaceBefore=6, spaceAfter=8)
    tom = ParagraphStyle("Tom", fontName=BOLD, fontSize=SIZE_TOM, alignment=TA_CENTER,
                         spaceBefore=0, spaceAfter=16, leading=SIZE_TOM * 1.2)
    czesc = ParagraphStyle("Czesc", fontName=BOLD, fontSize=SIZE_PART, alignment=TA_CENTER,
                           spaceBefore=0, spaceAfter=14, leading=SIZE_PART * 1.2)
    rozdz = ParagraphStyle("Rozdzial", fontName=BOLD, fontSize=SIZE_CHAPTER, alignment=TA_LEFT,
                           spaceBefore=6, spaceAfter=12, leading=SIZE_CHAPTER * 1.25)
    series = ParagraphStyle("series", fontName=REG, fontSize=SIZE_SERIES, alignment=TA_CENTER,
                            spaceAfter=6)
    bigt = ParagraphStyle("bigt", fontName=BOLD, fontSize=34, alignment=TA_CENTER, spaceAfter=18,
                          leading=34 * 1.2)
    blurb = ParagraphStyle("blurb", fontName=ITAL, fontSize=SIZE_BODY, alignment=TA_CENTER,
                           leading=SIZE_BODY * 1.4, leftIndent=1.2 * cm, rightIndent=1.2 * cm)
    tochdr = ParagraphStyle("tochdr", fontName=BOLD, fontSize=SIZE_PART, spaceAfter=14,
                            alignment=TA_CENTER)
    toc0 = ParagraphStyle("toc0", fontName=BOLD, fontSize=13, leading=20, spaceBefore=10)
    toc1 = ParagraphStyle("toc1", fontName=BOLD, fontSize=11, leading=16, leftIndent=14,
                          spaceBefore=4)
    toc2 = ParagraphStyle("toc2", fontName=REG, fontSize=10.5, leading=14, leftIndent=30)

    class IgielkiDoc(BaseDocTemplate):
        def __init__(self, path):
            super().__init__(path, pagesize=A4, title="Igiełki")
            self._h = 0
            fr = Frame(2.5 * cm, 2.2 * cm, A4[0] - 4.5 * cm, A4[1] - 4.4 * cm, id="b")
            self.addPageTemplates([PageTemplate(id="main", frames=[fr], onPage=self._foot)])

        def build(self, flowables, **kw):
            self._h = 0
            return super().build(flowables, **kw)

        def _foot(self, canvas, doc):
            canvas.saveState(); canvas.setFont(REG, 9)
            canvas.drawCentredString(A4[0] / 2, 1.2 * cm, str(doc.page))
            canvas.restoreState()

        def afterFlowable(self, flowable):
            if isinstance(flowable, Paragraph):
                sn = flowable.style.name
                if sn in ("Tom", "Czesc", "Rozdzial"):
                    text = flowable.getPlainText()
                    level = {"Tom": 0, "Czesc": 1, "Rozdzial": 2}[sn]
                    key = "h%d" % self._h; self._h += 1
                    self.canv.bookmarkPage(key)
                    self.canv.addOutlineEntry(text, key, level=level, closed=(level > 0))
                    self.notify("TOCEntry", (level, text, self.page, key))

    def make_toc():
        t = TableOfContents()
        t.levelStyles = [toc0, toc1, toc2]
        return t

    def chapter_flow(md, no):
        title = _title_of(md)
        head = f"Rozdział {no}. {title}" if no else title
        fl = [PageBreak(), Paragraph(_rml(head), rozdz)]
        for kind, text in body_lines(md):
            if kind == "scene":
                fl.append(Paragraph("✳", scene))
            elif kind == "quote":
                fl.append(Paragraph(_rml(text), quote))
            else:
                fl.append(Paragraph(_rml(text), body))
        return fl

    def part_flow(part_dir):
        fl = [PageBreak(), Paragraph(_rml(_title_of(part_dir / "_czesc.md")), czesc)]
        for ch in chapters_of(part_dir):
            fl += chapter_flow(ch, chapter_no(ch))
        return fl

    OUT_DIR.mkdir(exist_ok=True)
    made = []
    for kind, arg in targets:
        if kind == "single":
            tom_dir = arg
            n = tom_number(tom_dir)
            out = OUT_DIR / f"Igielki - Tom {n}.pdf"
            flow = [Spacer(1, 5 * cm), Paragraph(SERIES_TITLE, series),
                    Paragraph(_rml(_title_of(tom_dir / "_tom.md")), bigt),
                    PageBreak(), Paragraph("Spis treści", tochdr), make_toc(), PageBreak(),
                    Paragraph(_rml(_title_of(tom_dir / "_tom.md")), tom)]
            for part_dir in parts_of(tom_dir):
                flow += part_flow(part_dir)
            dod = tom_dir / "_dodatek.md"
            if dod.exists():
                flow += [PageBreak(), Paragraph(_rml(_title_of(dod)), czesc)]
                for kd, tx in body_lines(dod):
                    flow.append(Paragraph("✳", scene) if kd == "scene"
                                else Paragraph(_rml(tx), body))
            IgielkiDoc(str(out)).multiBuild(flow)
            print(f"Zapisano PDF: {out}"); made.append(out)
        else:
            toms = arg
            out = OUT_DIR / "Igielki - calosc.pdf"
            flow = [Spacer(1, 4 * cm), Paragraph(SERIES_TITLE, bigt),
                    Paragraph(SERIES_SUB, series), Spacer(1, 0.6 * cm),
                    Paragraph(SERIES_BLURB, blurb),
                    PageBreak(), Paragraph("Spis treści", tochdr), make_toc()]
            for tom_dir in toms:
                flow.append(PageBreak())
                flow.append(Paragraph(_rml(_title_of(tom_dir / "_tom.md")), tom))
                for part_dir in parts_of(tom_dir):
                    flow += part_flow(part_dir)
                dod = tom_dir / "_dodatek.md"
                if dod.exists():
                    flow += [PageBreak(), Paragraph(_rml(_title_of(dod)), czesc)]
                    for kd, tx in body_lines(dod):
                        flow.append(Paragraph("✳", scene) if kd == "scene"
                                    else Paragraph(_rml(tx), body))
            IgielkiDoc(str(out)).multiBuild(flow)
            print(f"Zapisano PDF: {out}"); made.append(out)
    return made


# ============================ DOCX (python-docx) ============================

def build_docx(targets):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _rfonts(rpr):
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts"); rpr.append(rf)
        for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rf.set(qn(a), FONT)

    def add_run(p, text, size, bold=False, italic=False):
        r = p.add_run(text); r.font.name = FONT; _rfonts(r._element.get_or_add_rPr())
        r.font.size = Pt(size); r.bold = bold; r.italic = italic
        lang = OxmlElement("w:lang"); lang.set(qn("w:val"), "pl-PL")
        r._element.get_or_add_rPr().append(lang)
        return r

    def add_inline(p, text, size, bold=False, italic=False):
        pos = 0
        for m in INLINE.finditer(text):
            if m.start() > pos:
                add_run(p, text[pos:m.start()], size, bold, italic)
            if m.group(1) is not None:
                add_run(p, m.group(1), size, True, italic)
            else:
                add_run(p, m.group(2), size, bold, True)
            pos = m.end()
        if pos < len(text):
            add_run(p, text[pos:], size, bold, italic)

    def setup():
        doc = Document()
        nrm = doc.styles["Normal"]; nrm.font.name = FONT; nrm.font.size = Pt(SIZE_BODY)
        nrm.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        nrm.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        nrm.paragraph_format.line_spacing = 1.35
        nrm.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        for name, size in (("Heading 1", SIZE_TOM), ("Heading 2", SIZE_PART),
                           ("Heading 3", SIZE_CHAPTER)):
            st = doc.styles[name]; st.font.name = FONT; st.font.size = Pt(size)
            st.font.bold = True; st.font.color.rgb = RGBColor(0, 0, 0)
            _rfonts(st.element.get_or_add_rPr()); st.paragraph_format.keep_with_next = True
        s = doc.sections[0]
        s.page_height = Cm(29.7); s.page_width = Cm(21.0)
        s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.2)
        s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)
        upd = OxmlElement("w:updateFields"); upd.set(qn("w:val"), "true")
        doc.settings.element.append(upd)
        fp = s.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = fp.add_run()
        for typ, txt in (("begin", None), ("instr", "PAGE"), ("end", None)):
            if typ == "instr":
                el = OxmlElement("w:instrText"); el.set(qn("xml:space"), "preserve"); el.text = txt
            else:
                el = OxmlElement("w:fldChar"); el.set(qn("w:fldCharType"), typ)
            rr._r.append(el)
        rr.font.name = FONT; rr.font.size = Pt(10)
        return doc

    def hding(doc, text, level, size, center=False, pb=False):
        p = doc.add_paragraph(style=f"Heading {level}")
        if center:
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if pb:
            p.paragraph_format.page_break_before = True
        add_inline(p, text, size, bold=True)

    def bodyp(doc, text, italic=False, center=False, indent=None):
        p = doc.add_paragraph(); add_inline(p, text, SIZE_BODY, italic=italic)
        if center:
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if indent:
            p.paragraph_format.left_indent = Cm(indent)

    def toc_field(doc):
        p = doc.add_paragraph(); run = p.add_run()
        b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
        ins = OxmlElement("w:instrText"); ins.set(qn("xml:space"), "preserve")
        ins.text = 'TOC \\o "1-3" \\h \\z \\u'
        sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
        tr = OxmlElement("w:r"); tt = OxmlElement("w:t")
        tt.text = "Spis treści — w Wordzie: Ctrl+A, potem F9."; tr.append(tt)
        e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
        run._r.append(b); run._r.append(ins); run._r.append(sep); p._p.append(tr)
        p.add_run()._r.append(e)

    def render_chapter(doc, md, no):
        title = _title_of(md)
        hding(doc, f"Rozdział {no}. {title}" if no else title, 3, SIZE_CHAPTER, pb=True)
        for kind, text in body_lines(md):
            if kind == "scene":
                bodyp(doc, "✳", center=True)
            elif kind == "quote":
                bodyp(doc, text, italic=True, indent=0.8)
            else:
                bodyp(doc, text)

    def render_tom(doc, tom_dir):
        hding(doc, _title_of(tom_dir / "_tom.md"), 1, SIZE_TOM, center=True, pb=True)
        for part_dir in parts_of(tom_dir):
            hding(doc, _title_of(part_dir / "_czesc.md"), 2, SIZE_PART, center=True, pb=True)
            for ch in chapters_of(part_dir):
                render_chapter(doc, ch, chapter_no(ch))
        dod = tom_dir / "_dodatek.md"
        if dod.exists():
            hding(doc, _title_of(dod), 2, SIZE_PART, center=True, pb=True)
            for kind, text in body_lines(dod):
                bodyp(doc, "✳", center=True) if kind == "scene" else bodyp(doc, text)

    OUT_DIR.mkdir(exist_ok=True); made = []
    for kind, arg in targets:
        doc = setup()
        if kind == "single":
            tom_dir = arg; n = tom_number(tom_dir)
            out = OUT_DIR / f"Igielki - Tom {n}.docx"
            for _ in range(5):
                doc.add_paragraph()
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, SERIES_TITLE, SIZE_SERIES)
            p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p2, _title_of(tom_dir / "_tom.md"), 28, bold=True)
            toc_field(doc)
            render_tom(doc, tom_dir)
        else:
            toms = arg; out = OUT_DIR / "Igielki - calosc.docx"
            for _ in range(5):
                doc.add_paragraph()
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, SERIES_TITLE, 34, bold=True)
            p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p2, SERIES_SUB, SIZE_SERIES)
            pb = doc.add_paragraph(); pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(pb, SERIES_BLURB, SIZE_BODY, italic=True)
            toc_field(doc)
            for tom_dir in toms:
                render_tom(doc, tom_dir)
        doc.save(str(out)); print(f"Zapisano DOCX: {out}"); made.append(out)
    return made


# ============================ CLI ============================

def resolve_targets(tom_arg, toms):
    if tom_arg is None:
        return [("all", toms)]
    t = str(tom_arg).strip().lower()
    if t == "wszystko":
        return [("single", d) for d in toms] + [("all", toms)]
    if t in ("calosc", "całość", "all"):
        return [("all", toms)]
    n = int(t)
    match = [d for d in toms if tom_number(d) == n]
    if not match:
        print(f"UWAGA: brak tomu {n}", file=sys.stderr); return []
    return [("single", match[0])]


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--typ", choices=["pdf", "docx", "oba"], default="pdf")
    ap.add_argument("--tom", default=None, help="N | calosc | wszystko (domyślnie: calosc)")
    ap.add_argument("--out", default=None)
    args = ap.parse_args()
    toms = list_toms()
    if not toms:
        print("Brak tomów w opowiadania/ — uruchom najpierw split_source.py", file=sys.stderr)
        return
    targets = resolve_targets(args.tom, toms)
    if not targets:
        return
    if args.typ in ("docx", "oba"):
        build_docx(targets)
    if args.typ in ("pdf", "oba"):
        build_pdf(targets)


if __name__ == "__main__":
    main()
